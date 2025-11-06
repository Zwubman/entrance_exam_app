import io
import os
import uuid
import zipfile
import textwrap
from PIL import Image
from docx import Document
from app.config.setting import settings

def extract_doc_data(file_bytes: bytes):
    file_stream = io.BytesIO(file_bytes)
    extracted_payloads = []
    text_segments = []
    image_paths = []

    try:
        # Try to process as DOCX first
        document = Document(file_stream)
        
        # Extract text from paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                sentences = textwrap.wrap(text, width=settings.CHUNK_LENGTH)
                text_segments.extend(sentences)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        sentences = textwrap.wrap(text, width=settings.CHUNK_LENGTH)
                        text_segments.extend(sentences)
        
        # Extract images
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx_zip:
                for file_name in docx_zip.namelist():
                    if file_name.startswith("word/media/"):
                        image_data = docx_zip.read(file_name)
                        try:
                            img = Image.open(io.BytesIO(image_data))
                            image_name = f"doc_image_{uuid.uuid4().hex}.png"
                            image_path = os.path.join(settings.UPLOADS_DIR, image_name)
                            img.save(image_path, "PNG")
                            image_paths.append(image_path)
                        except Exception as e:
                            print(f"Failed to extract image {file_name}: {e}")
        except Exception as e:
            print(f"Error reading images from DOCX: {e}")
            
    except Exception:
        # Fallback to textract for old .doc files
        try:
            import textract
            file_stream.seek(0)
            text = textract.process(file_stream, extension='doc').decode('utf-8')
            if text:
                sentences = textwrap.wrap(text.strip(), width=settings.CHUNK_LENGTH)
                text_segments.extend(sentences)
        except ImportError:
            print("textract not installed for .doc file support")
        except Exception as e:
            print(f"Error processing with textract: {e}")

    extracted_payloads.append({
        "texts": text_segments,
        "images": image_paths
    })

    return extracted_payloads